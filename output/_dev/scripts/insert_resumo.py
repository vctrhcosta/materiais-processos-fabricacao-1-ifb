#!/usr/bin/env python3
"""
insert_resumo.py — Insere RESUMO e ABSTRACT no DOCX existente,
substituindo os placeholders sem tocar no restante do documento.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from copy import deepcopy
import re

DOCX_PATH = Path(__file__).parent / "pesquisa-tintas-pigmentos-brasil.docx"

RESUMO_TEXT = (
    "O presente trabalho tem como objetivo analisar a indústria brasileira de tintas e pigmentos "
    "na perspectiva da indústria de transformação, identificando os nichos produtivos, as "
    "matérias-primas empregadas e os processos de fabricação característicos de cada segmento. "
    "A metodologia adotada consiste em pesquisa bibliográfica e documental, com base em dados "
    "setoriais da ABRAFATI, do Instituto Brasileiro de Geografia e Estatística (IBGE/PIA-Empresa), "
    "de relatórios de mercado da Grand View Research e da Mordor Intelligence, além de fontes "
    "técnicas especializadas. Os resultados indicam que o Brasil ocupa a 4ª posição no ranking "
    "mundial de produção de tintas, com 2,005 bilhões de litros produzidos em 2025, inserindo-se "
    "em um mercado global avaliado em USD 211,28 bilhões em 2024. O setor estrutura-se como elo "
    "intermediário entre as indústrias de base — petroquímica, mineração e química fina — e as "
    "indústrias usuárias, sendo classificado na Divisão 20 da CNAE 2.0. A formulação de tintas "
    "envolve quatro famílias de matérias-primas — resinas, pigmentos, solventes e aditivos — cujas "
    "combinações determinam os nichos produtivos identificados: imobiliário, industrial geral, "
    "automotivo, naval, tintas em pó, sistemas de cura UV/EB e tintas especiais e funcionais. Os "
    "processos de fabricação variam conforme o nicho, com destaque para a dispersão de pigmentos "
    "em moinhos de pérolas, a extrusão para tintas em pó e a fotopolimerização para sistemas de "
    "cura por radiação. Conclui-se que o setor apresenta elevado potencial de diversificação, com "
    "nichos estratégicos de alta oportunidade em anticorrosivos para o segmento offshore, tintas "
    "intumescentes, revestimentos para energias renováveis e sistemas base água, tendência "
    "impulsionada por regulamentações ambientais e pela crescente demanda ESG."
)

PALAVRAS_CHAVE = "tintas industriais; pigmentos; processos de fabricação; indústria de transformação; nichos produtivos."

ABSTRACT_TEXT = (
    "This paper aims to analyze the Brazilian paint and pigment industry from the perspective of "
    "the manufacturing industry, identifying productive niches, raw materials, and manufacturing "
    "processes characteristic of each segment. The methodology consists of bibliographic and "
    "documentary research, based on sectoral data from ABRAFATI, the Brazilian Institute of "
    "Geography and Statistics (IBGE/PIA-Empresa), market reports from Grand View Research and "
    "Mordor Intelligence, and specialized technical sources. The results indicate that Brazil "
    "ranks 4th in the global paint production ranking, with 2.005 billion liters produced in 2025, "
    "within a global market valued at USD 211.28 billion in 2024. The sector is structured as an "
    "intermediate link between base industries — petrochemical, mining, and fine chemistry — and "
    "end-user industries, classified under Division 20 of Brazil's National Classification of "
    "Economic Activities (CNAE 2.0). Paint formulation involves four families of raw materials — "
    "resins, pigments, solvents, and additives — whose combinations define the productive niches "
    "identified: residential, general industrial, automotive, naval, powder coatings, UV/EB curing "
    "systems, and specialty and functional paints. Manufacturing processes vary by niche, with "
    "emphasis on pigment dispersion in bead mills, extrusion for powder coatings, and "
    "photopolymerization for radiation-curing systems. It is concluded that the sector presents "
    "high diversification potential, with strategic high-opportunity niches in anticorrosive "
    "coatings for the offshore segment, intumescent paints, renewable energy coatings, and "
    "waterborne systems, a trend driven by environmental regulations and growing ESG demand."
)

KEYWORDS = "industrial paints; pigments; manufacturing processes; manufacturing industry; productive niches."

# Marcadores de localização nos placeholders gerados pelo normalize_md.py
RESUMO_MARKER   = "RESUMO"
ABSTRACT_MARKER = "ABSTRACT"
PLACEHOLDER_RESUMO   = "Parágrafo único com"
PLACEHOLDER_ABSTRACT = "Same content as Resumo"


def fmt_body(para):
    """Aplica formatação padrão ABNT ao parágrafo de corpo."""
    para.style.font.name = "Times New Roman"
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_run(para, text, bold=False, font_size=12):
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return run


def find_paragraph_index(doc, marker):
    """Retorna o índice do parágrafo que contém o marcador."""
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            return i
    return None


def replace_placeholder_block(doc, title_marker, body_marker,
                               section_title, body_text, keywords_label, keywords_text):
    """
    Localiza o bloco:
      TITULO (ex: RESUMO)
      [placeholder body]
      Palavras-chave: [...]
    e substitui pelo conteúdo real.
    """
    paragraphs = doc.paragraphs

    # Encontra o título
    title_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip().startswith(title_marker) and len(p.text.strip()) < 60:
            title_idx = i
            break

    if title_idx is None:
        print(f"  ⚠ Marcador '{title_marker}' não encontrado.")
        return False

    # Encontra o parágrafo placeholder do corpo (próximos 10 parágrafos)
    body_idx = None
    for i in range(title_idx + 1, min(title_idx + 10, len(paragraphs))):
        if body_marker in paragraphs[i].text:
            body_idx = i
            break

    if body_idx is None:
        print(f"  ⚠ Placeholder de corpo '{body_marker}' não encontrado após '{title_marker}'.")
        return False

    # Encontra o parágrafo de palavras-chave (próximos 5 após o body)
    kw_idx = None
    for i in range(body_idx + 1, min(body_idx + 5, len(paragraphs))):
        if "palavra" in paragraphs[i].text.lower() or "keyword" in paragraphs[i].text.lower():
            kw_idx = i
            break

    # Substitui o placeholder do corpo pelo texto real
    body_para = paragraphs[body_idx]
    body_para.clear()
    fmt_body(body_para)
    set_run(body_para, body_text)

    # Substitui o placeholder das palavras-chave
    if kw_idx is not None:
        kw_para = paragraphs[kw_idx]
        kw_para.clear()
        fmt_body(kw_para)
        kw_para.paragraph_format.first_line_indent = Cm(0)
        set_run(kw_para, f"{keywords_label} ", bold=True)
        set_run(kw_para, keywords_text)

    print(f"  ✓ {section_title} inserido (parágrafo {body_idx})")
    if kw_idx:
        print(f"  ✓ {keywords_label} inserido (parágrafo {kw_idx})")
    return True


def main():
    if not DOCX_PATH.exists():
        print(f"Erro: arquivo não encontrado — {DOCX_PATH}")
        return

    doc = Document(str(DOCX_PATH))
    print(f"Abrindo: {DOCX_PATH.name} ({len(doc.paragraphs)} parágrafos)")

    replace_placeholder_block(
        doc,
        title_marker=RESUMO_MARKER,
        body_marker=PLACEHOLDER_RESUMO,
        section_title="RESUMO",
        body_text=RESUMO_TEXT,
        keywords_label="Palavras-chave:",
        keywords_text=PALAVRAS_CHAVE,
    )

    replace_placeholder_block(
        doc,
        title_marker=ABSTRACT_MARKER,
        body_marker=PLACEHOLDER_ABSTRACT,
        section_title="ABSTRACT",
        body_text=ABSTRACT_TEXT,
        keywords_label="Keywords:",
        keywords_text=KEYWORDS,
    )

    doc.save(str(DOCX_PATH))
    print(f"\n✓ Salvo: {DOCX_PATH.name}")


if __name__ == "__main__":
    main()
