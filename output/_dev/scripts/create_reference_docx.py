#!/usr/bin/env python3
"""
create_reference_docx.py — Cria reference.docx com estilos ABNT NBR 14724/2024
Para uso com pandoc: pandoc input.md --reference-doc=reference.docx -o output.docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_number_to_header(section):
    """Adiciona número de página no canto superior direito (fonte 10, 2cm da borda)."""
    header = section.header
    # Garante que existe parágrafo no header
    if not header.paragraphs:
        header.add_paragraph()
    para = header.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    run = para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    # Campo PAGE do Word
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


def configure_style_normal(styles):
    """Corpo do texto: TNR 12, justificado, 1,5 entrelinhas, recuo 1,25cm."""
    s = styles['Normal']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.font.bold = False
    s.font.italic = False
    pf = s.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def configure_heading(styles, level: int, bold: bool, italic: bool, uppercase_xml: bool = False):
    """Configura estilo de título conforme hierarquia ABNT."""
    style_name = f'Heading {level}'
    s = styles[style_name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.font.bold = bold
    s.font.italic = italic
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.font.underline = False
    pf = s.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Remove qualquer cor de fundo / shading herdado
    rPr = s.element.find(qn('w:rPr'))
    if rPr is None:
        from docx.oxml import OxmlElement as OE
        rPr = OE('w:rPr')
        s.element.append(rPr)

    # Heading 1: numeração de página quebra e MAIÚSCULAS via XML
    if level == 1:
        pPr = s.element.get_or_add_pPr()
        # Page break before section
        pbr = OxmlElement('w:pageBreakBefore')
        pbr.set(qn('w:val'), '1')
        pPr.append(pbr)
        # Keep with next
        kwn = OxmlElement('w:keepWithNext')
        kwn.set(qn('w:val'), '1')
        pPr.append(kwn)
        # Caps (maiúsculas) — aplica via allCaps no rPr do estilo
        if uppercase_xml:
            caps = OxmlElement('w:caps')
            caps.set(qn('w:val'), '1')
            rPr.append(caps)


def configure_block_text(styles):
    """Citação longa (>3 linhas): TNR 10, simples, recuo 4cm da margem esquerda."""
    try:
        s = styles['Block Text']
    except KeyError:
        s = styles.add_style('Block Text', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0, 0, 0)
    pf = s.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.left_indent = Cm(4.0)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)


def configure_caption_style(styles):
    """Caption para Quadros/Figuras/Tabelas: TNR 10, centralizado, simples."""
    try:
        s = styles['Caption']
    except KeyError:
        s = styles.add_style('Caption', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.bold = False
    s.font.italic = False
    s.font.color.rgb = RGBColor(0, 0, 0)
    pf = s.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)


def configure_source_style(styles):
    """Fonte abaixo de Quadros/Tabelas: TNR 10, alinhado à esquerda, simples."""
    try:
        s = styles['Source']
    except KeyError:
        s = styles.add_style('Source', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    s.font.bold = False
    s.font.color.rgb = RGBColor(0, 0, 0)
    pf = s.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)


def configure_references_style(styles):
    """Referências: TNR 12, esquerda, simples, sem recuo."""
    try:
        s = styles['Bibliography']
    except KeyError:
        s = styles.add_style('Bibliography', 1)
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(0, 0, 0)
    pf = s.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)  # espaço simples entre referências


def configure_footnote_style(styles):
    """Notas de rodapé: TNR 10, simples."""
    try:
        s = styles['Footnote Text']
    except KeyError:
        return
    s.font.name = 'Times New Roman'
    s.font.size = Pt(10)
    pf = s.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)


def create_abnt_reference_docx(output_path: str):
    doc = Document()
    styles = doc.styles

    # ── Configuração da página (A4, margens ABNT anverso) ─────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)    # Anverso: esquerda 3cm
    section.top_margin = Cm(3.0)     # Anverso: superior 3cm
    section.right_margin = Cm(2.0)   # Anverso: direita 2cm
    section.bottom_margin = Cm(2.0)  # Anverso: inferior 2cm
    section.header_distance = Cm(1.25)  # Número de página a ~2cm da borda superior
    section.footer_distance = Cm(1.5)

    # ── Número de página no cabeçalho ─────────────────────────────────────────
    add_page_number_to_header(section)

    # ── Estilos ───────────────────────────────────────────────────────────────
    configure_style_normal(styles)

    # Seção primária:   MAIÚSCULA, negrito, 12
    configure_heading(styles, 1, bold=True, italic=False, uppercase_xml=True)
    # Seção secundária: caixa baixa, negrito, 12
    configure_heading(styles, 2, bold=True, italic=False)
    # Seção terciária:  caixa baixa, itálico, negrito, 12
    configure_heading(styles, 3, bold=True, italic=True)
    # Seção quaternária: caixa baixa, itálico sem negrito, 12
    configure_heading(styles, 4, bold=False, italic=True)
    # Seção quinária:   caixa baixa, sem negrito, sem itálico, 12
    try:
        configure_heading(styles, 5, bold=False, italic=False)
    except Exception:
        pass

    configure_block_text(styles)
    configure_caption_style(styles)
    configure_source_style(styles)
    configure_references_style(styles)
    configure_footnote_style(styles)

    # ── Placeholder mínimo para pandoc não reclamar de doc vazio ──────────────
    doc.add_paragraph('', style='Normal')

    doc.save(output_path)
    print(f"✓ reference.docx criado: {output_path}")


if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 else 'reference.docx'
    create_abnt_reference_docx(out)
